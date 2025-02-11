from datetime import datetime
from docx import Document
import openpyxl
import json

def process_dados():
    alunos = []
    try:
        with open('alunos.json', 'r', encoding='utf-8') as arquivo:
            return json.load(arquivo)
    except FileNotFoundError:
        return []

def save_dados(alunos):
    with open('alunos.json', 'w', encoding='utf-8') as arquivo:
        json.dump(alunos, arquivo, indent=4, ensure_ascii=False)

def add_alun(alunos):
    nome = input('Digite o nome do(a) aluno(a): ')
    matricula = int(input('Digite o número de matrícula do(a) aluno(a): '))
    birth = input('Digite a data de nascimento do(a) aluno(a) (no formato DD/MM/AAAA): ')
    try:
        data = datetime.strptime(birth, '%d/%m/%Y')
        birth_format = data.strftime('%d-%m-%Y')
    except ValueError:
        print('Formato de data inválido... Por favor, tente novamente da forma indicada.')
        return
    notas = []
    
    qtd = int(input('Quantas notas deseja adicionar ao perfil do(a) aluno(a): '))
    for i in range(qtd):
        nota = float(input(f'Digite a {i+1}° nota: '))
        notas.append(nota)
    media = sum(notas) / len(notas)

    while True:
        try:
            faltas = int(input('Quantas faltas o(a) aluno(a) apresentou: '))
            break
        except ValueError:
            print('Valor inválido, tente novamente.')

    alunos.append({'nome': nome, 'matricula': matricula, 'data_nascimento': birth_format, 'notas': notas, 'media': media, 'faltas': faltas})
    save_dados(alunos)
    print('\nAluno(a) adicionado ao sistema.\n')

def calcular_status(media, faltas):
    if media >= 7 and faltas <= 5:
        return "Aprovado"
    elif media >= 5 and faltas <= 5:
        return "Recuperação"
    else:
        return "Reprovado"

def list_al(alunos):
    if not alunos:
        print('\nSem alunos registrados no sistema.\n')
    else:
        for aluno in alunos:
            print(f'Nome: {aluno["nome"]}\nMatrícula: {aluno["matricula"]}\nData de nascimento: {aluno["data_nascimento"]}\nNotas: {aluno["notas"]}\nMédia: {aluno["media"]}\nFaltas: {aluno["faltas"]}\nMédia: {aluno["media"]}')

def find_al(alunos):
    while True:
        try:
            filtro = int(input('Digite o número de matrícula do aluno: '))
            break
        except ValueError:
            print('Valor não válido, tente novamente.')
    for aluno in alunos:
        if filtro == aluno['matricula']:
            print(f'\nAluno(a) encontrado:\nNome: {aluno["nome"]}\nMatrícula: {aluno["matricula"]}\nData de nascimento: {aluno["data_nascimento"]}\nNotas: {aluno["notas"]}\nMédia: {aluno["media"]}\nFaltas: {aluno["faltas"]}\nMédia: {aluno["media"]}\n')
            return aluno
    print('\nAluno não encontrado...')
    return None

def relatorio_individual(alunos):
    aluno = find_al(alunos)
    if aluno:
        doc = Document()
        doc.add_heading(f'Relatório de {aluno["nome"]}', level=1)
        doc.add_paragraph(f'Matrícula: {aluno["matricula"]}')
        doc.add_paragraph(f'Data de Nascimento: {aluno["data_nascimento"]}')
        doc.add_paragraph(f'Notas: {aluno["notas"]}')
        doc.add_paragraph(f'Média: {aluno["media"]:.2f}')
        doc.add_paragraph(f'Faltas: {aluno["faltas"]}')
        doc.add_paragraph(f'Status: {calcular_status(aluno["media"], aluno["faltas"])}')
        doc.save(f'relatorio_{aluno["matricula"]}.docx')
        print(f'Relatório gerado: relatorio_{aluno["matricula"]}.docx\n')

def modificar(alunos):
    aluno = find_al(alunos)
    if aluno:
        print(f'O que gostaria de alterar:\n1 - Nome\n2 - Matrícula\n3 - Data de nascimento\n4 - Notas\n5 - Faltas\n6 - Cancelar')
        opc = int(input('Digite o número da opção que deseja: '))
        if opc == 1:
            aluno['nome'] = input('Digite o novo nome: ')
        elif opc == 2:
            while True:
                try:
                    matricula = int(input('Digite o novo número de matrícula: '))
                    aluno['matricula'] = matricula
                    break
                except ValueError:
                    print('Valor inválido, tente novamente.')
        elif opc == 3:
            while True:
                birth = input('Digite a nova data de nascimento(no formato DD/MM/AAAA): ')
                try:
                    data = datetime.strptime(birth, '%d/%m/%Y')
                    aluno['data_nascimento'] = data.strftime('%d-%m-%Y')
                    break
                except ValueError:
                    print('Formato de data inválido, tente novamente.')
        elif opc == 4:
            try:
                qtd = int(input('Quantas notas deseja alterar: '))
                aluno['notas'] = []
                for i in range(qtd):
                    print(f"As notas atuais são: {aluno['notas']}")
                    new_nota = float(input(f'Digite a {i+1}° nova nota: '))
                    aluno['notas'].append(new_nota)
            except ValueError:
                print('Erro, valor de nota inválido.')
        elif opc == 5:
            while True:
                try:
                    faltas = int(input('Digite o novo número de faltas do(a) aluno(a): '))
                    aluno['faltas'] = faltas
                    break
                except ValueError:
                    print('Valor inválido, tente novamente.')
        elif opc == 6:
            print('Voltando ao menu inicial...')
            return
        else:
            print('\nOpção inválida, tente novamente.\n')
        save_dados(alunos)
        print('\nDados atualizados.\n')

def excluir(alunos):
    aluno = find_al(alunos)
    if aluno:
        alunos.remove(aluno)
        save_dados(alunos)
        print('\nRemoção realizada.\n')

def organizar(alunos):
    filtro = input('- Organizar por nome\n- Organizar por número de matrícula\n- Organizar por data de nascimento\nDigite: "nome" ou "matricula" ou "data de nascimento": ')
    if filtro.lower() == 'nome':
        alunos.sort(key=lambda aluno: aluno['nome'])
    elif filtro.lower() == 'matricula':
        alunos.sort(key=lambda aluno: aluno['matricula'])
    elif filtro.lower() == 'data de nascimento':
        alunos.sort(key=lambda aluno: aluno['data_nascimento'])
    else:
        print('\nOpção inválida\n')
        return
    
    save_dados(alunos)
    print(f'Alunos ordenados por {filtro}')

def relatorio_grupo(alunos):
    doc = Document()
    doc.add_heading('Relatório de Alunos', level=1)
    for aluno in alunos:
        doc.add_paragraph(f'Nome: {aluno["nome"]}')
        doc.add_paragraph(f'Matrícula: {aluno["matricula"]}')
        doc.add_paragraph(f'Data de Nascimento: {aluno["data_nascimento"]}')
        doc.add_paragraph(f'Notas: {aluno["notas"]}')
        doc.add_paragraph(f'Média: {aluno["media"]:.2f}')
        doc.add_paragraph(f'Faltas: {aluno["faltas"]}')
        doc.add_paragraph(f'Status: {calcular_status(aluno["media"], aluno["faltas"])}')
        doc.add_paragraph('-' * 40)
    doc.save('relatorio_grupo.docx')
    print('Relatório de grupo gerado: relatorio_grupo.docx')

def relatorio_excel(alunos):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Nome", "Matrícula", "Data de Nascimento", "Notas", "Média", "Faltas", "Status"])
    for aluno in alunos:
        status = calcular_status(aluno["media"], aluno["faltas"])
        ws.append([
            aluno["nome"], aluno["matricula"], aluno["data_nascimento"], 
            ', '.join(map(str, aluno["notas"])), aluno["media"], aluno["faltas"], status
        ])
    wb.save('relatorio_alunos.xlsx')
    print('Exportação para Excel concluída: relatorio_alunos.xlsx')

def menu():
    alunos = process_dados()
    while True:
        print(f'{"=" * 50}')
        print(f'{"SISTEMA DE GESTÃO DE ALUNOS".center(50)}')
        print(f'{"=" * 50}')
        print('\n1 - Adicionar um novo aluno\n2 - Listar alunos já registrados\n3 - Buscar aluno\n4 - Editar dados de um aluno já cadastrado\n5 - Remover um aluno do sistema\n6 - Organizar lista de alunos\n7 - Relatório de um aluno em particular\n8 - Relatório de turma geral\n9 - Planilha de desempenho dos alunos\n0 - Sair do sistema')
        decisao = int(input('- '))
        if decisao == 1:
            add_alun(alunos)
        elif decisao == 2:
            list_al(alunos)
        elif decisao == 3:
            find_al(alunos)
        elif decisao == 4:
            modificar(alunos)
        elif decisao == 5:
            excluir(alunos)
        elif decisao == 6:
            organizar(alunos)
        elif decisao == 7:
            relatorio_individual(alunos)
        elif decisao == 8:
            relatorio_grupo(alunos)
        elif decisao == 9:
            relatorio_excel(alunos)
        elif decisao == 0:
            print('Saindo do sistema...')
            break
        else:
            print('Opção inválida, tente novamente.')

menu()
